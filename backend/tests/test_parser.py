import os
import sys
import tempfile
import unittest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document as DocxDocument

from app.ingestion.parser import UnsupportedFileType, parse_document


class TestParseDocx(unittest.TestCase):
    def _make_docx(self, paragraphs_per_section, num_sections):
        """Tạo file .docx thật để test — không phải mock."""
        doc = DocxDocument()
        for s in range(num_sections):
            for p in range(paragraphs_per_section):
                doc.add_paragraph(f"Đoạn {s}-{p}: nội dung ví dụ về RAG và LLM.")
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)
        return tmp.name

    def test_parses_docx_into_sections_grouped_by_paragraph_count(self):
        path = self._make_docx(paragraphs_per_section=3, num_sections=2)
        try:
            sections = parse_document(path, paragraphs_per_section=3)
            self.assertEqual(len(sections), 2)
            position_ref, text = sections[0]
            self.assertTrue(position_ref.startswith("Mục"))
            self.assertIn("Đoạn 0-0", text)
            self.assertIn("Đoạn 0-1", text)
            self.assertIn("Đoạn 0-2", text)
        finally:
            os.remove(path)

    def test_empty_paragraphs_are_skipped(self):
        doc = DocxDocument()
        doc.add_paragraph("")
        doc.add_paragraph("Nội dung thật duy nhất.")
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)
        try:
            sections = parse_document(tmp.name, paragraphs_per_section=5)
            joined = " ".join(text for _, text in sections)
            self.assertIn("Nội dung thật duy nhất.", joined)
        finally:
            os.remove(tmp.name)

    def test_unsupported_extension_raises(self):
        tmp = tempfile.NamedTemporaryFile(suffix=".txt", delete=False)
        tmp.write(b"noi dung")
        tmp.close()
        try:
            with self.assertRaises(UnsupportedFileType):
                parse_document(tmp.name)
        finally:
            os.remove(tmp.name)

    def test_missing_file_raises_file_not_found(self):
        with self.assertRaises(FileNotFoundError):
            parse_document("/khong/ton/tai.pdf")


if __name__ == "__main__":
    unittest.main()
