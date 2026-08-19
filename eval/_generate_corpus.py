"""Sinh corpus tài liệu DOCX thật (nội dung giáo dục chính xác, tự viết) cho
eval/golden_set.jsonl. Chạy 1 lần để tạo file trong eval/documents/, không
phải một phần của bộ test tự động của backend.
"""

import os

from docx import Document

OUT_DIR = os.path.join(os.path.dirname(__file__), "documents")
os.makedirs(OUT_DIR, exist_ok=True)


def write_doc(filename, sections):
    doc = Document()
    for heading, paragraphs in sections:
        doc.add_heading(heading, level=1)
        for p in paragraphs:
            doc.add_paragraph(p)
    doc.save(os.path.join(OUT_DIR, filename))
    print("Wrote", filename)


write_doc(
    "ML_Optimization.docx",
    [
        (
            "Gradient Descent",
            [
                "Gradient Descent là một thuật toán tối ưu lặp, dùng để tìm giá trị tham số "
                "làm cực tiểu hoá một hàm mất mát (loss function). Ở mỗi bước, thuật toán "
                "tính đạo hàm (gradient) của hàm mất mát theo từng tham số, rồi cập nhật "
                "tham số theo hướng ngược với gradient — vì gradient chỉ hướng tăng nhanh "
                "nhất của hàm số, nên đi ngược hướng đó sẽ làm hàm số giảm.",
                "Công thức cập nhật cơ bản: tham số_mới = tham số_cũ - learning_rate * gradient. "
                "Quá trình này lặp lại nhiều vòng (epoch) cho tới khi hàm mất mát hội tụ về "
                "một giá trị đủ nhỏ hoặc không giảm thêm đáng kể.",
            ],
        ),
        (
            "Learning Rate",
            [
                "Learning rate là siêu tham số quyết định độ lớn của mỗi bước cập nhật "
                "trong Gradient Descent. Nếu learning rate quá lớn, các bước cập nhật có "
                "thể nhảy qua điểm cực tiểu và khiến hàm mất mát dao động hoặc phân kỳ "
                "(không hội tụ). Nếu learning rate quá nhỏ, quá trình hội tụ sẽ rất chậm, "
                "tốn nhiều epoch để đạt kết quả tốt.",
                "Trong thực hành, learning rate thường được chọn trong khoảng 0.001 đến 0.1 "
                "tuỳ bài toán, và có thể giảm dần theo thời gian huấn luyện (learning rate "
                "schedule) để vừa hội tụ nhanh ở giai đoạn đầu, vừa ổn định ở giai đoạn cuối.",
            ],
        ),
        (
            "Các biến thể của Gradient Descent",
            [
                "Batch Gradient Descent tính gradient trên TOÀN BỘ tập dữ liệu huấn luyện "
                "trước khi cập nhật tham số một lần — chính xác nhưng chậm với dữ liệu lớn.",
                "Stochastic Gradient Descent (SGD) cập nhật tham số sau MỖI một mẫu dữ liệu "
                "— nhanh hơn nhưng đường đi tới cực tiểu dao động nhiều hơn.",
                "Mini-batch Gradient Descent là dạng dung hoà phổ biến nhất hiện nay: chia dữ "
                "liệu thành các batch nhỏ (vd 32, 64, 128 mẫu) và cập nhật tham số sau mỗi batch.",
            ],
        ),
    ],
)

write_doc(
    "ML_DecisionTree.docx",
    [
        (
            "Decision Tree là gì",
            [
                "Decision Tree (cây quyết định) là một mô hình học máy có giám sát, dùng "
                "được cho cả bài toán phân loại (classification) lẫn hồi quy (regression). "
                "Mô hình biểu diễn quá trình ra quyết định dưới dạng cây, trong đó mỗi nút "
                "trong (internal node) là một điều kiện kiểm tra trên một thuộc tính dữ "
                "liệu, mỗi nhánh là một kết quả của điều kiện đó, và mỗi nút lá (leaf node) "
                "là một nhãn dự đoán.",
            ],
        ),
        (
            "Cách chọn thuộc tính để phân nhánh",
            [
                "Ở mỗi bước xây cây, thuật toán chọn thuộc tính giúp CHIA dữ liệu thành các "
                "nhóm càng thuần nhất (pure) càng tốt. Hai độ đo phổ biến để đánh giá độ "
                "thuần nhất là Gini Impurity và Entropy (dùng trong Information Gain). "
                "Thuộc tính nào giúp giảm Gini/Entropy nhiều nhất sau khi chia sẽ được chọn "
                "làm điều kiện phân nhánh tại nút đó.",
                "Quá trình chia nhánh lặp lại đệ quy cho tới khi đạt điều kiện dừng, ví dụ: "
                "nút đã thuần nhất hoàn toàn, đạt độ sâu tối đa cho phép, hoặc số mẫu trong "
                "nút quá ít để chia tiếp.",
            ],
        ),
        (
            "Ưu và nhược điểm",
            [
                "Ưu điểm: dễ diễn giải (interpretable), không cần chuẩn hoá dữ liệu đầu vào, "
                "xử lý được cả thuộc tính dạng số lẫn dạng phân loại.",
                "Nhược điểm: dễ bị overfitting nếu cây quá sâu (học thuộc lòng dữ liệu huấn "
                "luyện), và có thể không ổn định — thay đổi nhỏ trong dữ liệu huấn luyện có "
                "thể tạo ra một cây rất khác. Random Forest (tổng hợp nhiều cây) thường được "
                "dùng để khắc phục nhược điểm này.",
            ],
        ),
    ],
)

write_doc(
    "DL_CNN.docx",
    [
        (
            "CNN là gì",
            [
                "Convolutional Neural Network (CNN) là một kiến trúc mạng nơ-ron chuyên "
                "dùng để xử lý dữ liệu có cấu trúc lưới, phổ biến nhất là ảnh. CNN khai thác "
                "tính cục bộ (locality) và tính bất biến theo vị trí (translation invariance) "
                "của ảnh để giảm đáng kể số lượng tham số so với mạng nơ-ron kết nối đầy đủ "
                "(fully-connected) truyền thống.",
            ],
        ),
        (
            "Phép tích chập (Convolution)",
            [
                "Lớp tích chập (convolutional layer) áp dụng một bộ lọc (kernel/filter) kích "
                "thước nhỏ (vd 3x3) trượt qua toàn bộ ảnh đầu vào, tại mỗi vị trí tính tích "
                "vô hướng giữa kernel và vùng ảnh tương ứng để tạo ra một giá trị trong "
                "feature map đầu ra. Mỗi kernel học được để phát hiện một loại đặc trưng cụ "
                "thể, ví dụ cạnh ngang, cạnh dọc, hoặc kết cấu (texture).",
                "Stride là bước nhảy của kernel khi trượt qua ảnh; padding là kỹ thuật thêm "
                "viền (thường là 0) quanh ảnh đầu vào để kiểm soát kích thước feature map "
                "đầu ra và tránh mất thông tin ở biên ảnh.",
            ],
        ),
        (
            "Pooling",
            [
                "Lớp pooling (thường là Max Pooling) làm giảm kích thước không gian của "
                "feature map bằng cách lấy giá trị lớn nhất (hoặc trung bình, với Average "
                "Pooling) trong từng vùng nhỏ. Việc này giúp giảm số lượng tham số, tăng tốc "
                "tính toán, và làm mô hình bền vững hơn trước các dịch chuyển nhỏ của đối "
                "tượng trong ảnh.",
            ],
        ),
    ],
)

write_doc(
    "DL_NeuralNetwork.docx",
    [
        (
            "Neural Network cơ bản",
            [
                "Mạng nơ-ron nhân tạo gồm nhiều lớp (layer) các nơ-ron kết nối với nhau, mỗi "
                "kết nối có một trọng số (weight). Mỗi nơ-ron nhận đầu vào có trọng số, cộng "
                "thêm một hệ số bias, rồi đưa qua một hàm kích hoạt phi tuyến (activation "
                "function, vd ReLU, Sigmoid) để tạo đầu ra truyền tới lớp tiếp theo.",
            ],
        ),
        (
            "Backpropagation",
            [
                "Backpropagation (lan truyền ngược) là thuật toán dùng để tính gradient của "
                "hàm mất mát theo TỪNG trọng số trong mạng, để phục vụ bước cập nhật trọng "
                "số bằng Gradient Descent. Thuật toán áp dụng quy tắc chuỗi (chain rule) của "
                "đạo hàm, lan truyền lỗi từ lớp đầu ra ngược về lớp đầu vào, tính gradient "
                "của từng lớp dựa trên gradient đã tính được ở lớp phía sau nó.",
                "Nhờ backpropagation, việc tính gradient cho một mạng có hàng triệu tham số "
                "trở nên khả thi về mặt tính toán, thay vì phải tính đạo hàm trực tiếp cho "
                "từng tham số một cách độc lập.",
            ],
        ),
        (
            "Vanishing Gradient",
            [
                "Vanishing Gradient là hiện tượng gradient trở nên CỰC KỲ NHỎ khi lan truyền "
                "ngược qua nhiều lớp trong mạng sâu, khiến các lớp đầu (gần input) gần như "
                "không được cập nhật trọng số, làm mạng học rất chậm hoặc không học được. "
                "Hiện tượng này thường xảy ra khi dùng hàm kích hoạt Sigmoid/Tanh ở mạng quá "
                "sâu, vì đạo hàm của các hàm này luôn nhỏ hơn 1, nhân dồn qua nhiều lớp làm "
                "gradient tiến về 0.",
                "Các kỹ thuật khắc phục phổ biến gồm: dùng hàm kích hoạt ReLU (đạo hàm không "
                "bị chặn trên và bằng 1 với đầu vào dương), khởi tạo trọng số hợp lý, và dùng "
                "kiến trúc có kết nối tắt (skip connection) như trong ResNet.",
            ],
        ),
    ],
)

print("Done. Corpus generated in", OUT_DIR)
